"""Extract text from DOCX or TXT resume files."""

from pathlib import Path

from docx import Document


def extract_text_from_docx(path: str | Path) -> str:
    """
    Extract text from a DOCX or TXT file.
    - DOCX: paragraphs and table cells via python-docx.
    - TXT: read file as plain text.
    Raises FileNotFoundError if path does not exist.
    Returns empty string on invalid DOCX.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix != ".docx":
        return path.read_text(encoding="utf-8", errors="replace")

    try:
        doc = Document(path)
    except Exception:
        return ""

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n\n".join(parts) if parts else ""
